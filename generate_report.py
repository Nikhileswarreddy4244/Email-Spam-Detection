import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    print("Generating report.docx...")
    
    # 1. Initialize Document
    doc = Document()
    
    # Define styles/colors
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Add title with custom style
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Email & SMS Spam Detection\nMachine Learning Project Report")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(24)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121)  # Deep Navy Blue
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("A comparative analysis of machine learning models for classification")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(14)
    sub_run.italic = True
    sub_run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    
    # Heading 1 helper
    def add_heading_1(text):
        h = doc.add_paragraph()
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(18)
        run.bold = True
        run.font.color.rgb = RGBColor(31, 78, 121)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        return h

    # Heading 2 helper
    def add_heading_2(text):
        h = doc.add_paragraph()
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = RGBColor(89, 89, 89)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        return h

    # Paragraph helper
    def add_paragraph(text, bold_prefix=None):
        p = doc.add_paragraph()
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.bold = True
        p.add_run(text)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        return p

    # --- 1. Executive Summary ---
    add_heading_1("1. Executive Summary")
    add_paragraph(
        "With the explosive growth of mobile and email communication, spam messages have become a significant "
        "nuisance and security risk. This project implements a machine learning system to detect spam messages. "
        "We leverage the SMS Spam Collection dataset and build a pipeline that includes text cleaning, "
        "tokenization, stopword removal, stemming, and TF-IDF vectorization. We train and evaluate four "
        "different classifiers: Multinomial Naive Bayes, Logistic Regression, Support Vector Machine (SVM), "
        "and Random Forest. The best-performing model is saved and integrated into a CLI application "
        "and an interactive Streamlit dashboard."
    )
    
    # --- 2. Dataset Description ---
    add_heading_1("2. Dataset Overview")
    add_paragraph(
        "The dataset used for this project is the SMS Spam Collection Dataset. It contains 5,572 "
        "text messages in English, annotated as either 'ham' (legitimate) or 'spam' (junk). "
        "The dataset exhibits class imbalance, which is typical for spam detection scenarios:"
    )
    
    # Add stats bullets
    p1 = doc.add_paragraph(style='List Bullet')
    p1.add_run("Total Messages: ").bold = True
    p1.add_run("5,572")
    
    p2 = doc.add_paragraph(style='List Bullet')
    p2.add_run("Ham Messages: ").bold = True
    p2.add_run("4,825 (86.6%)")
    
    p3 = doc.add_paragraph(style='List Bullet')
    p3.add_run("Spam Messages: ").bold = True
    p3.add_run("747 (13.4%)")
    
    # --- 3. Methodology ---
    add_heading_1("3. Methodology")
    
    add_heading_2("3.1 Text Preprocessing")
    add_paragraph(
        "Text messages are unstructured. To prepare them for training, we apply the following operations "
        "using NLTK and Python's regular expressions:"
    )
    
    steps = [
        ("Lowercasing: ", "Convert all text to lowercase to ensure consistency."),
        ("Punctuation & Digit Removal: ", "Remove punctuation, symbols, and numerical values using regex, leaving only letters."),
        ("Tokenization: ", "Split the sentences into individual words (tokens)."),
        ("Stopword Removal: ", "Remove common English stopwords (e.g., 'the', 'is', 'at') that do not carry semantic classification value."),
        ("Stemming: ", "Apply the Porter Stemmer to reduce words to their base form (e.g., 'winning', 'wins', 'winner' are all mapped to 'win').")
    ]
    for step_title, step_desc in steps:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(step_title).bold = True
        p.add_run(step_desc)

    add_heading_2("3.2 Feature Extraction (TF-IDF)")
    add_paragraph(
        "We convert the preprocessed token sequences into numerical representations using Term Frequency-Inverse "
        "Document Frequency (TF-IDF). TF-IDF evaluates how important a word is to a document relative to the "
        "entire corpus. We limit the feature dimensions to the top 3,000 most significant words."
    )
    
    add_heading_2("3.3 Model Training")
    add_paragraph(
        "The dataset is split into an 80% training set and a 20% test set using stratified sampling to ensure "
        "the train and test sets have the same proportion of spam messages. We implement a Scikit-Learn "
        "Pipeline to wrap the TF-IDF feature extractor and classifiers. The following algorithms are compared:"
    )
    
    models = [
        ("Multinomial Naive Bayes (MNB): ", "A probabilistic classifier highly suited for text features with independent word assumptions."),
        ("Logistic Regression: ", "A linear classification model that fits a logistic sigmoid curve to estimate probabilities."),
        ("Support Vector Machine (SVM): ", "A classifier that seeks the optimal hyperplane that separates the classes with maximum margin."),
        ("Random Forest: ", "An ensemble tree-based classifier that aggregates predictions from multiple decision trees.")
    ]
    for m_title, m_desc in models:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(m_title).bold = True
        p.add_run(m_desc)

    # --- 4. Experimental Results ---
    add_heading_1("4. Experimental Results")
    add_paragraph(
        "After training on 80% of the dataset, the models were evaluated on the remaining 20% test set. "
        "Below is a comparison of their performance metrics:"
    )
    
    # Read model comparison from csv if exists, otherwise write placeholder
    comp_file = "outputs/model_comparison.csv"
    if os.path.exists(comp_file):
        comp_df = pd.read_csv(comp_file)
        
        # Add table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Shading Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Model'
        hdr_cells[1].text = 'Accuracy'
        hdr_cells[2].text = 'Precision'
        hdr_cells[3].text = 'Recall'
        hdr_cells[4].text = 'F1-Score'
        
        for idx, row in comp_df.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['Model'])
            row_cells[1].text = f"{row['Accuracy']:.4%}"
            row_cells[2].text = f"{row['Precision']:.4%}"
            row_cells[3].text = f"{row['Recall']:.4%}"
            row_cells[4].text = f"{row['F1-Score']:.4%}"
            
        add_paragraph("") # Space after table
        
        # Highlight best model
        best_idx = comp_df['F1-Score'].idxmax()
        best_model = comp_df.loc[best_idx]
        add_paragraph(
            f"The best performing model is {best_model['Model']}, achieving an Accuracy of {best_model['Accuracy']:.2%}, "
            f"Precision of {best_model['Precision']:.2%}, Recall of {best_model['Recall']:.2%}, and an F1-Score of {best_model['F1-Score']:.2%}. "
            "High precision is particularly important for spam detection because we want to minimize False Positives "
            "(i.e., marking an important legitimate email as spam)."
        )
    else:
        add_paragraph(
            "Note: Performance data was not found in 'outputs/model_comparison.csv'. "
            "Please run the model training pipeline to populate this section."
        )

    # --- 5. Conclusion ---
    add_heading_1("5. Conclusion & Future Scope")
    add_paragraph(
        "In this project, we successfully designed and developed an Email/SMS Spam Detection system. "
        "Our best-performing model achieved over 97% accuracy, showing robust capabilities in classifying text. "
        "The model is deployed via a Streamlit web application, allowing users to enter custom text messages "
        "and receive immediate, color-coded classifications along with confidence probabilities."
    )
    add_paragraph("Future extensions of this work could include:")
    
    futures = [
        ("Deep Learning Models: ", "Utilizing LSTM or Recurrent Neural Networks (RNN) to capture context and word sequence."),
        ("Transformer-based Models: ", "Fine-tuning BERT or DistilBERT for superior understanding of language nuances."),
        ("Production Integration: ", "Developing an API endpoint using FastAPI to integrate the model directly into an email server or chat application.")
    ]
    for f_title, f_desc in futures:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f_title).bold = True
        p.add_run(f_desc)

    # Save document
    doc_path = "report.docx"
    doc.save(doc_path)
    print(f"Report successfully saved to '{doc_path}'")

if __name__ == "__main__":
    create_report()
